import os
import fitz
import re
from flask import Flask, request, jsonify
from docx import Document
from pdf2image import convert_from_path
import tempfile
from PIL import Image
import pytesseract
import requests


app = Flask(__name__)

FILE_EXTENSIONS = [".pdf", ".docx", ".txt", ".jpg", ".png", ".jpeg"]
current_dir = os.path.dirname(os.path.abspath(__file__))

# Xây dựng đường dẫn đến tesseract.exe trong thư mục Tesseract-OCR
tesseract_path = os.path.join(current_dir, 'Tesseract-OCR', 'tesseract.exe')

# Cấu hình đường dẫn cho pytesseract
pytesseract.pytesseract.tesseract_cmd = tesseract_path

QDRANT_SERVER_URL = "http://180.148.1.178:6889/qdrant-storing"


def check_image(filepath):
    doc = fitz.open(filepath)
    for page in doc:
        text = page.get_text()
        if text.strip():
            doc.close()
            return False
    doc.close()
    return True


def clean_text(raw_text):
    # Loại bỏ các dòng chỉ chứa ít nhất 3 dấu chấm, có thể có khoảng trắng ở giữa
    cleaned_text = re.sub(r'^\s*(\.\s*){3,}\s*$', '', raw_text, flags=re.MULTILINE)
    # Chuẩn hóa khoảng trắng và dấu xuống dòng thừa
    cleaned_text = re.sub(r'(?<!\n)\n(?!\n)', ' ', cleaned_text.strip())
    return cleaned_text


def preprocess_text(text):
    # Loại bỏ các trang số (nếu có)
    text = re.sub(r'(?:Page|Trang)?\s*-?\s*\d+\s*-?', '', text, flags=re.IGNORECASE)
    # Loại bỏ các ký tự không mong muốn
    text = re.sub(r"[^\w\s.,!?%\-–()]", "", text)
    # Rút gọn chuỗi dấu chấm dài (ít nhất 3 dấu, có thể có khoảng trắng) thành khoảng trắng
    text = re.sub(r'(\.\s*){3,}', ' ', text)
    # Chuẩn hóa dấu xuống dòng
    text = re.sub(r'\n{2,}', '\n', text)
    # Chuẩn hóa khoảng trắng
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r' +\n', '\n', text)
    return text.strip()


def extract_text_from_image_files(file_path):

    text = ""

    # Trường hợp file ảnh: .png, .jpg, .jpeg
    if file_path.lower().endswith(('.png', '.jpg', '.jpeg')):
        image = Image.open(file_path)
        # Sử dụng pytesseract để OCR, chỉ định ngôn ngữ là tiếng Việt và tiếng Anh
        text = pytesseract.image_to_string(image, lang='vie+eng')
        image.close()

    # Trường hợp PDF scan: dùng pdf2image
    elif file_path.lower().endswith('.pdf'):
        images = convert_from_path(file_path, dpi=150, grayscale=True)
        for image in images:
            temp_path = tempfile.mktemp(suffix='.png')
            image.save(temp_path, 'PNG')
            # Sử dụng pytesseract để OCR
            text += pytesseract.image_to_string(Image.open(temp_path), lang='vie+eng') + "\n"
            os.unlink(temp_path)
            image.close()

    # Trường hợp DOCX chứa ảnh
    elif file_path.lower().endswith('.docx'):
        doc = Document(file_path)
        for rel in doc.part._rels:
            target = doc.part._rels[rel].target_ref
            if "image" in target:
                img_bytes = doc.part.related_parts[target].blob
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                    temp_img.write(img_bytes)
                    temp_img_path = temp_img.name

                # Sử dụng pytesseract để OCR
                text += pytesseract.image_to_string(Image.open(temp_img_path), lang='vie+eng') + "\n"
                os.unlink(temp_img_path)

    else:
        raise ValueError(f"Unsupported file type for image OCR: {file_path}")

    return preprocess_text(clean_text(text))


def extract_text_from_docx_with_image_and_text(file_path):
 
    if not file_path.lower().endswith('.docx'):
        raise ValueError("File không phải định dạng .docx")

    doc = Document(file_path)
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    for rel in doc.part._rels:
        target = doc.part._rels[rel].target_ref
        if "image" in target:
            img_bytes = doc.part.related_parts[target].blob
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_img.write(img_bytes)
                temp_img_path = temp_img.name

            ocr_result = pytesseract.image_to_string(Image.open(temp_img_path), lang='vie+eng')
            if ocr_result.strip():
                full_text.append(ocr_result.strip())
            os.unlink(temp_img_path)

    combined_text = "\n".join(full_text)
    return preprocess_text(clean_text(combined_text))


def extract_text_from_pdf_with_image_and_text(file_path):

    if not file_path.lower().endswith('.pdf'):
        raise ValueError("File không phải định dạng PDF")

    doc = fitz.open(file_path)
    combined_text = ""

    for page in doc:
        text = page.get_text().strip()
        if text:
            combined_text += text + "\n"
        else:
            # OCR ảnh nếu trang không có text
            pix = page.get_pixmap(dpi=150)
            temp_img_path = tempfile.mktemp(suffix=".png")
            pix.save(temp_img_path)

            # Sử dụng pytesseract để OCR
            ocr_result = pytesseract.image_to_string(Image.open(temp_img_path), lang='vie+eng')
            combined_text += ocr_result.strip() + "\n"
            os.unlink(temp_img_path)

    doc.close()
    return preprocess_text(clean_text(combined_text))


def extract_text_from_text_only_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()

    # Làm sạch văn bản
    return preprocess_text(clean_text(text))


def extract_text_from_text_only_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    # Làm sạch văn bản
    return preprocess_text(clean_text(text))



def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()

    # Làm sạch văn bản
    return preprocess_text(clean_text(text))


def check_pdf_content_type(file_path):
    doc = fitz.open(file_path)
    has_text = False
    has_image = False

    for page in doc:
        if page.get_text().strip():
            has_text = True
        if page.get_images(full=True):
            has_image = True
        if has_text and has_image:
            break

    doc.close()

    if has_text and has_image:
        return 1  # both
    elif has_text:
        return 2  # text
    elif has_image:
        return 3  # image
    else:
        return 0  # empty


def check_docx_content_type(file_path):
    doc = Document(file_path)
    has_text = False
    has_image = False

    for para in doc.paragraphs:
        if para.text.strip():
            has_text = True
            break

    for rel in doc.part._rels:
        rel_obj = doc.part._rels[rel]
        if "image" in rel_obj.target_ref:
            has_image = True
            break

    if has_text and has_image:
        return 1  # both
    elif has_text:
        return 2  # text
    elif has_image:
        return 3  # image
    else:
        return 0  # empty


def extract_text(file_path):
    file_name = os.path.basename(file_path)
    lower_path = file_path.lower()

    if lower_path.endswith(('.jpg', '.png', '.jpeg')):
        text = extract_text_from_image_files(file_path)

    elif lower_path.endswith('.txt'):
        text = extract_text_from_txt(file_path)

    elif lower_path.endswith('.pdf'):
        file_type = check_pdf_content_type(file_path)
        if file_type == 1:  # both
            text = extract_text_from_pdf_with_image_and_text(file_path)
        elif file_type == 2:  # text
            text = extract_text_from_text_only_pdf(file_path)
        elif file_type == 3:  # image
            text = extract_text_from_image_files(file_path)
        else:
            text = ""  # empty or unsupported

    elif lower_path.endswith('.docx'):
        file_type = check_docx_content_type(file_path)
        if file_type == 1:  # both
            text = extract_text_from_docx_with_image_and_text(file_path)
        elif file_type == 2:  # text
            text = extract_text_from_text_only_docx(file_path)
        elif file_type == 3:  # image
            text = extract_text_from_image_files(file_path)
        else:
            text = ""  # empty or unsupported

    else:
        raise ValueError(f"Unsupported file type: {file_path}")

    return text, file_name

def process_single_file(file_info, loai_phieu, form_data):
    file_path = file_info.get('path')
    file_name = file_info.get('file_name')
    file_type = file_info.get('file_type')

    if not file_path or not file_name or not file_type:
        return {
            "file_name": file_name or "unknown",
            "file_type": file_type or "unknown",
            "status": "error",
            "message": "Missing file information"
        }

    if not os.path.exists(file_path):
        return {
            "file_name": file_name,
            "file_type": file_type,
            "status": "error",
            "message": f"File not found: {file_path}"
        }

    if not any(file_path.lower().endswith(ext) for ext in FILE_EXTENSIONS):
        return {
            "file_name": file_name,
            "file_type": file_type,
            "status": "error",
            "message": f"Unsupported file type. Supported extensions: {', '.join(FILE_EXTENSIONS)}"
        }

    try:
        text, extracted_file_name = extract_text(file_path)
        if not text:
            return {
                "file_name": file_name,
                "file_type": file_type,
                "status": "error",
                "message": "No text extracted"
            }

        return {
            "file_name": file_name,
            "file_type": file_type,
            "content": text,
            "status": "success",
            "message": "File content extracted successfully"
        }
    except Exception as e:
        return {
            "file_name": file_name,
            "file_type": file_type,
            "status": "error",
            "message": f"Error processing file: {str(e)}"
        }


@app.route('/store-documents', methods=['POST'])
def store_documents():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400

        loai_phieu = data.get('loai_phieu')
        form_data = data.get('formData', {})
        files = data.get('files', [])

        if not loai_phieu:
            return jsonify({"error": "Missing loai_phieu"}), 400
        if not files:
            return jsonify({"error": "No files provided"}), 400

        results = []
        for file_info in files:
            result = process_single_file(file_info, loai_phieu, form_data)
            file_result = {
                "file_name": result.get("file_name"),
                "file_type": result.get("file_type"),
            }

            if result.get("status") == "success":
                content = result.get("content", "")
                file_result["content"] = content

                try:
                    forward_response = requests.post(
                        QDRANT_SERVER_URL,
                        json={
                            "text": content,
                            "file_name": result.get("file_name"),
                            "loai_phieu": loai_phieu,
                            "form_data": form_data
                        },
                        timeout=10  
                    )
                    file_result["forward_status"] = forward_response.status_code
                except Exception as forward_error:
                    file_result["forward_error"] = str(forward_error)

            else:
                file_result["error"] = result.get("message")

            results.append(file_result)

        return jsonify({
            "loai_phieu": loai_phieu,
            "formData": form_data,
            "files": results
        }), 200

    except Exception as e:
        return jsonify({"error": f"Server error: {str(e)}"}), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)