import os
import fitz
import re
from flask import Flask, request, jsonify
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from qdrant_client import QdrantClient
from qdrant_client.http.models import PointStruct, VectorParams, Distance
import uuid
from docx import Document
from pdf2image import convert_from_path
import tempfile
from langchain_experimental.text_splitter import SemanticChunker
from PIL import Image
import pytesseract


app = Flask(__name__)

FILE_EXTENSIONS = [".pdf", ".docx", ".txt", ".jpg", ".png", ".jpeg"]
GENERAL_COLLECTION_NAME = "general_documents"
current_dir = os.path.dirname(os.path.abspath(__file__))

# Xây dựng đường dẫn đến tesseract.exe trong thư mục Tesseract-OCR
tesseract_path = os.path.join(current_dir, 'Tesseract-OCR', 'tesseract.exe')

# Cấu hình đường dẫn cho pytesseract
pytesseract.pytesseract.tesseract_cmd = tesseract_path
embed_model = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2",
    model_kwargs={"device": "cpu"}
)
client = QdrantClient(url="http://localhost:6333")
created_collections = set()  # Cache for collection existence


def check_image(filepath):
    doc = fitz.open(filepath)
    for page in doc:
        text = page.get_text()
        if text.strip():
            doc.close()
            return False
    doc.close()
    return True


def clean_text(raw_text):
    # Loại bỏ các dòng chỉ chứa ít nhất 3 dấu chấm, có thể có khoảng trắng ở giữa
    cleaned_text = re.sub(r'^\s*(\.\s*){3,}\s*$', '', raw_text, flags=re.MULTILINE)
    # Chuẩn hóa khoảng trắng và dấu xuống dòng thừa
    cleaned_text = re.sub(r'(?<!\n)\n(?!\n)', ' ', cleaned_text.strip())
    return cleaned_text


def preprocess_text(text):
    # Loại bỏ các trang số (nếu có)
    text = re.sub(r'(?:Page|Trang)?\s*-?\s*\d+\s*-?', '', text, flags=re.IGNORECASE)
    # Loại bỏ các ký tự không mong muốn
    text = re.sub(r"[^\w\s.,!?%\-–()]", "", text)
    # Rút gọn chuỗi dấu chấm dài (ít nhất 3 dấu, có thể có khoảng trắng) thành khoảng trắng
    text = re.sub(r'(\.\s*){3,}', ' ', text)
    # Chuẩn hóa dấu xuống dòng
    text = re.sub(r'\n{2,}', '\n', text)
    # Chuẩn hóa khoảng trắng
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r' +\n', '\n', text)
    return text.strip()


def extract_text_from_image_files(file_path):
    """
    Hàm này xử lý các file chỉ chứa ảnh (.png, .jpg, .jpeg, .pdf scan, .docx chứa ảnh)
    và trả về văn bản đã OCR bằng pytesseract.
    """
    text = ""

    # Trường hợp file ảnh: .png, .jpg, .jpeg
    if file_path.lower().endswith(('.png', '.jpg', '.jpeg')):
        image = Image.open(file_path)
        # Sử dụng pytesseract để OCR, chỉ định ngôn ngữ là tiếng Việt và tiếng Anh
        text = pytesseract.image_to_string(image, lang='vie+eng')
        image.close()

    # Trường hợp PDF scan: dùng pdf2image
    elif file_path.lower().endswith('.pdf'):
        images = convert_from_path(file_path, dpi=150, grayscale=True)
        for image in images:
            temp_path = tempfile.mktemp(suffix='.png')
            image.save(temp_path, 'PNG')
            # Sử dụng pytesseract để OCR
            text += pytesseract.image_to_string(Image.open(temp_path), lang='vie+eng') + "\n"
            os.unlink(temp_path)
            image.close()

    # Trường hợp DOCX chứa ảnh
    elif file_path.lower().endswith('.docx'):
        doc = Document(file_path)
        for rel in doc.part._rels:
            target = doc.part._rels[rel].target_ref
            if "image" in target:
                img_bytes = doc.part.related_parts[target].blob
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                    temp_img.write(img_bytes)
                    temp_img_path = temp_img.name

                # Sử dụng pytesseract để OCR
                text += pytesseract.image_to_string(Image.open(temp_img_path), lang='vie+eng') + "\n"
                os.unlink(temp_img_path)

    else:
        raise ValueError(f"Unsupported file type for image OCR: {file_path}")

    return preprocess_text(clean_text(text))


def extract_text_from_docx_with_image_and_text(file_path):
    """
    Trích xuất nội dung từ file .docx có cả text và hình ảnh (OCR).
    Kết hợp cả đoạn văn bản và OCR từ ảnh trong file.
    """
    if not file_path.lower().endswith('.docx'):
        raise ValueError("File không phải định dạng .docx")

    doc = Document(file_path)
    full_text = []

    # 1. Text từ đoạn văn
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # 2. Text từ ảnh trong docx
    for rel in doc.part._rels:
        target = doc.part._rels[rel].target_ref
        if "image" in target:
            img_bytes = doc.part.related_parts[target].blob
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_img.write(img_bytes)
                temp_img_path = temp_img.name

            # OCR với pytesseract
            ocr_result = pytesseract.image_to_string(Image.open(temp_img_path), lang='vie+eng')
            if ocr_result.strip():
                full_text.append(ocr_result.strip())
            os.unlink(temp_img_path)

    combined_text = "\n".join(full_text)
    return preprocess_text(clean_text(combined_text))


def extract_text_from_pdf_with_image_and_text(file_path):
    """
    Trích xuất nội dung từ PDF có cả text và ảnh (OCR nếu cần).
    Nếu trang có text → lấy text; nếu không có → OCR ảnh.
    """
    if not file_path.lower().endswith('.pdf'):
        raise ValueError("File không phải định dạng PDF")

    doc = fitz.open(file_path)
    combined_text = ""

    for page in doc:
        text = page.get_text().strip()
        if text:
            combined_text += text + "\n"
        else:
            # OCR ảnh nếu trang không có text
            pix = page.get_pixmap(dpi=150)
            temp_img_path = tempfile.mktemp(suffix=".png")
            pix.save(temp_img_path)

            # Sử dụng pytesseract để OCR
            ocr_result = pytesseract.image_to_string(Image.open(temp_img_path), lang='vie+eng')
            combined_text += ocr_result.strip() + "\n"
            os.unlink(temp_img_path)

    doc.close()
    return preprocess_text(clean_text(combined_text))


def extract_text_from_text_only_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()

    # Làm sạch văn bản
    return preprocess_text(clean_text(text))


def extract_text_from_text_only_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    # Làm sạch văn bản
    return preprocess_text(clean_text(text))



def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()

    # Làm sạch văn bản
    return preprocess_text(clean_text(text))


def check_pdf_content_type(file_path):
    doc = fitz.open(file_path)
    has_text = False
    has_image = False

    for page in doc:
        if page.get_text().strip():
            has_text = True
        if page.get_images(full=True):
            has_image = True
        if has_text and has_image:
            break

    doc.close()

    if has_text and has_image:
        return 1  # both
    elif has_text:
        return 2  # text
    elif has_image:
        return 3  # image
    else:
        return 0  # empty


def check_docx_content_type(file_path):
    doc = Document(file_path)
    has_text = False
    has_image = False

    for para in doc.paragraphs:
        if para.text.strip():
            has_text = True
            break

    for rel in doc.part._rels:
        rel_obj = doc.part._rels[rel]
        if "image" in rel_obj.target_ref:
            has_image = True
            break

    if has_text and has_image:
        return 1  # both
    elif has_text:
        return 2  # text
    elif has_image:
        return 3  # image
    else:
        return 0  # empty


def extract_text(file_path):
    file_name = os.path.basename(file_path)
    lower_path = file_path.lower()

    if lower_path.endswith(('.jpg', '.png', '.jpeg')):
        text = extract_text_from_image_files(file_path)

    elif lower_path.endswith('.txt'):
        text = extract_text_from_txt(file_path)

    elif lower_path.endswith('.pdf'):
        file_type = check_pdf_content_type(file_path)
        if file_type == 1:  # both
            text = extract_text_from_pdf_with_image_and_text(file_path)
        elif file_type == 2:  # text
            text = extract_text_from_text_only_pdf(file_path)
        elif file_type == 3:  # image
            text = extract_text_from_image_files(file_path)
        else:
            text = ""  # empty or unsupported

    elif lower_path.endswith('.docx'):
        file_type = check_docx_content_type(file_path)
        if file_type == 1:  # both
            text = extract_text_from_docx_with_image_and_text(file_path)
        elif file_type == 2:  # text
            text = extract_text_from_text_only_docx(file_path)
        elif file_type == 3:  # image
            text = extract_text_from_image_files(file_path)
        else:
            text = ""  # empty or unsupported

    else:
        raise ValueError(f"Unsupported file type: {file_path}")

    return text, file_name


def filter_invalid_chunks(chunks, min_length=30):
    filtered = []
    for chunk in chunks:
        cleaned = chunk.strip()
        if len(cleaned) >= min_length and not re.fullmatch(r"[.?!,:;\"']+", cleaned):
            filtered.append(cleaned)
    return filtered

def semantic_chunking(text, embed_model):
    try:
        # Khởi tạo SemanticChunker để chia văn bản thành các chunk ngữ nghĩa
        semantic_splitter = SemanticChunker(
            embeddings=embed_model,
            breakpoint_threshold_type="percentile",
            breakpoint_threshold_amount=95
        )
        # Chia văn bản thành các chunk ngữ nghĩa
        semantic_chunks = semantic_splitter.split_text(text)
        
        # Trả về tất cả các chunk, không lọc chunk ngắn
        return semantic_chunks
    except Exception as e:
        print(f"Error during chunking: {e}")
        return []


def ensure_collection(client, collection_name, embed_model):
    if collection_name not in created_collections:
        try:
            client.get_collection(collection_name)
            created_collections.add(collection_name)
        except:
            # Get embedding dimension by embedding a dummy text
            embedding = embed_model.embed_query("test")
            client.create_collection(
                collection_name=collection_name,
                vectors_config=VectorParams(
                    size=len(embedding),
                    distance=Distance.COSINE
                )
            )
            created_collections.add(collection_name)


def store_in_qdrant(chunks, file_name, collection_name, form_data, embed_model, client):
    batch_size = 16
    embeddings = []
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i:i + batch_size]
        embeddings.extend(embed_model.embed_documents(batch))

    points = [
        PointStruct(
            id=str(uuid.uuid4()),
            vector=embedding,
            payload={
                "file_name": file_name,
                "chunk_id": i,
                "text": chunk,
                **(form_data or {})
            }
        )
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings))
    ]
    client.upsert(collection_name=collection_name, points=points)


def process_single_file(file_info, loai_phieu, form_data, embed_model, client):
    file_path = file_info.get('path')
    file_name = file_info.get('file_name')
    file_type = file_info.get('file_type')

    if not file_path or not file_name or not file_type:
        return {
            "file_name": file_name or "unknown",
            "status": "error",
            "message": "Missing file information"
        }

    if not os.path.exists(file_path):
        return {
            "file_name": file_name,
            "status": "error",
            "message": f"File not found: {file_path}"
        }

    if not any(file_path.lower().endswith(ext) for ext in FILE_EXTENSIONS):
        return {
            "file_name": file_name,
            "status": "error",
            "message": f"Unsupported file type. Supported extensions: {', '.join(FILE_EXTENSIONS)}"
        }

    try:
        text, extracted_file_name = extract_text(file_path)
        if not text:
            return {
                "file_name": file_name,
                "status": "error",
                "message": "No text extracted"
            }

        chunks = semantic_chunking(text, embed_model)
        if not chunks:
            return {
                "file_name": file_name,
                "status": "error",
                "message": "No chunks created"
            }

        ensure_collection(client, loai_phieu, embed_model)
        ensure_collection(client, GENERAL_COLLECTION_NAME, embed_model)
        store_in_qdrant(chunks, file_name, loai_phieu, form_data, embed_model, client)
        store_in_qdrant(chunks, file_name, GENERAL_COLLECTION_NAME, form_data, embed_model, client)

        return {
            "file_name": file_name,
            "status": "success",
            "message": f"Processed: {len(chunks)} chunks created and stored",
            "content": chunks
        }
    except Exception as e:
        return {
            "file_name": file_name,
            "status": "error",
            "message": f"Error processing file: {str(e)}"
        }


@app.route('/store-documents', methods=['POST'])
def store_documents():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400

        loai_phieu = data.get('loai_phieu')
        form_data = data.get('formData', {})
        files = data.get('files', [])

        if not loai_phieu:
            return jsonify({"error": "Missing loai_phieu"}), 400
        if not files:
            return jsonify({"error": "No files provided"}), 400

        results = []
        for file_info in files:
            result = process_single_file(file_info, loai_phieu, form_data, embed_model, client)
            results.append(result)

        return jsonify({
            "status": "completed",
            "results": results
        }), 200

    except Exception as e:
        return jsonify({"error": f"Server error: {str(e)}"}), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)