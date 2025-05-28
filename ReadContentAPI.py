import os
import re
import fitz
from flask import Flask, request, jsonify
from pdf2image import convert_from_path
import easyocr
from docx import Document
import tabula
import tempfile
import pandas as pd

app = Flask(__name__)

FILE_EXTENSIONS = [".pdf", ".docx", ".txt", ".jpg", ".png", ".jpeg"]

def check_image(filepath):
    """Check if a PDF contains only images (no extractable text)."""
    try:
        doc = fitz.open(filepath)
        for page in doc:
            text = page.get_text().strip()
            if text:
                doc.close()
                return False
        doc.close()
        return True
    except Exception as e:
        print(f"Error checking image PDF: {e}")
        return True

def clean_text(raw_text):
    """Clean raw text by normalizing whitespace and line breaks."""
    text = re.sub(r'(?<!\n)\n(?!\n)', ' ', raw_text.strip())
    text = re.sub(r'\n{2,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r' +\n', '\n', text)
    return text.strip()

def preprocess_text(text):
    """Preprocess text by removing page numbers and unnecessary characters."""
    text = re.sub(r'(?:Page|Trang)?\s*-?\s*\d+\s*-?', '', text, flags=re.IGNORECASE)
    text = re.sub(r"[^\w\s.,!?%\-–()]", "", text)
    return clean_text(text)

def extract_text_with_ocr(file_path):
    """Extract text from images or image-based PDFs using EasyOCR."""
    try:
        reader = easyocr.Reader(['vi', 'en'], gpu=False)  # Disable GPU for broader compatibility
        text = ""
        if file_path.lower().endswith('.pdf'):
            images = convert_from_path(file_path, poppler_path=None)  # Ensure poppler is in PATH
            for i, image in enumerate(images):
                temp_path = tempfile.mktemp(suffix='.png')
                image.save(temp_path, 'PNG')
                results = reader.readtext(temp_path, detail=0, paragraph=True)  # Group text into paragraphs
                text += f"--- Page {i + 1} ---\n" + "\n".join(results) + "\n\n"
                os.unlink(temp_path)
        elif file_path.lower().endswith(('.png', '.jpg', '.jpeg')):
            results = reader.readtext(file_path, detail=0, paragraph=True)
            text += "\n".join(results)
        return preprocess_text(text)
    except Exception as e:
        print(f"OCR extraction failed: {e}")
        return ""

def extract_tables_from_pdf(file_path):
    """Extract tables from PDF using tabula-py with improved formatting."""
    try:
        # Try lattice first for structured tables, fallback to stream
        tables = tabula.read_pdf(file_path, pages='all', multiple_tables=True, lattice=True, stream=True)
        table_texts = []
        for i, table in enumerate(tables):
            # Handle missing values and ensure consistent formatting
            table = table.fillna('')
            # Convert table to markdown-like format for readability
            table_text = f"### Table {i + 1}\n"
            table_text += table.to_markdown(index=False, tablefmt="grid") + "\n\n"
            table_texts.append(preprocess_text(table_text))
        return "\n".join(table_texts)
    except Exception as e:
        print(f"PDF table extraction failed: {e}")
        return ""

def extract_tables_from_docx(file_path):
    """Extract tables from DOCX with improved formatting."""
    try:
        doc = Document(file_path)
        table_texts = []
        for i, table in enumerate(doc.tables):
            table_data = []
            # Extract column headers if available
            headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            # Format as markdown-like table
            table_text = f"### Table {i + 1}\n"
            if headers:
                table_text += "| " + " | ".join(headers) + " |\n"
                table_text += "| " + " | ".join(["---"] * len(headers)) + " |\n"
            table_text += "\n".join(["| " + " | ".join(row) + " |" for row in table_data]) + "\n\n"
            table_texts.append(preprocess_text(table_text))
        return "\n".join(table_texts)
    except Exception as e:
        print(f"DOCX table extraction failed: {e}")
        return ""

def extract_text_and_tables(file_path):
    """Extract text and tables from a file based on its type."""
    file_name = os.path.basename(file_path)
    text = ""
    tables = ""

    if not os.path.exists(file_path):
        return {"error": f"File not found: {file_path}"}, file_name

    if not any(file_path.lower().endswith(ext) for ext in FILE_EXTENSIONS):
        return {"error": f"Unsupported file type: {file_path}. Supported extensions: {', '.join(FILE_EXTENSIONS)}"}, file_name

    try:
        if file_path.lower().endswith(('.jpg', '.png', '.jpeg')) or (
                file_path.lower().endswith('.pdf') and check_image(file_path)):
            text = extract_text_with_ocr(file_path)
        elif file_path.lower().endswith('.pdf'):
            doc = fitz.open(file_path)
            text_blocks = []
            for page_num, page in enumerate(doc, 1):
                blocks = page.get_text("blocks")  # Extract text as blocks to preserve structure
                page_text = []
                for block in blocks:
                    block_text = block[4].strip()  # Block[4] is the text content
                    if block_text:
                        page_text.append(block_text)
                if page_text:
                    text_blocks.append(f"--- Page {page_num} ---\n" + "\n".join(page_text))
            text = "\n\n".join(text_blocks)
            doc.close()
            text = preprocess_text(clean_text(text))
            tables = extract_tables_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            text = preprocess_text(clean_text(text))
            tables = extract_tables_from_docx(file_path)
        elif file_path.lower().endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            text = preprocess_text(clean_text(text))

        combined_content = f"{text}\n\n{tables}".strip() if tables else text
        return {"content": combined_content, "file_name": file_name}, file_name

    except Exception as e:
        return {"error": f"Error processing file {file_path}: {str(e)}"}, file_name

@app.route('/process_file', methods=['POST'])
def process_file():
    """Flask API endpoint to process a file and return extracted content."""
    data = request.get_json()
    if not data or 'file_path' not in data:
        return jsonify({"error": "Missing 'file_path' in request body"}), 400

    file_path = data['file_path']
    result, file_name = extract_text_and_tables(file_path)
    if "error" in result:
        return jsonify(result), 400

    content = result["content"]
    if not content:
        return jsonify({"error": f"No content extracted from {file_path}"}), 400

    return jsonify({
        "file_name": file_name,
        "content": content
    }), 200

if __name__ == "__main__":
    app.run(debug=True, host='0.0.0.0', port=5001)